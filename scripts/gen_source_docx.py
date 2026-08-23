"""
生成软件著作权源代码鉴别材料 Word 文档
- 前30页 + 后30页，每页50行
- 页眉：软件名称 + 版本号
- 页码
"""
import os
import glob
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SOFTWARE_NAME = "One Work 智能协作平台软件"
VERSION = "V2.1"
LINES_PER_PAGE = 50
PAGES = 30
PROJECT_ROOT = r"D:\aionui-m0\1oneUI"
OUTPUT = r"D:\aionui-m0\1oneUI\One_Work_源代码鉴别材料.docx"

# dream 命名空间替换映射（按长度从长到短排序，避免部分替换）
NAME_REPLACEMENTS = [
    ("@aionui/", "@dream/"),
    ("AIONUI_", "DREAM_"),
    ("aioncore", "dreamcore"),
    ("aionrs", "dreamrs"),
    ("AionUi", "Dream"),
    ("AIONUI", "DREAM"),
    ("aionui", "dream"),
    ("AION_", "DREAM_"),
    ("1ONE Work", "One Work"),
    ("1ONE Code", "One Work"),
    ("1ONE", "One Work"),
]


def collect_source_files(root):
    files = []
    for ext in ("*.ts", "*.tsx"):
        files.extend(glob.glob(os.path.join(root, "packages", "**", ext), recursive=True))
    files = [f for f in files if "node_modules" not in f and "\\out\\" not in f and "\\dist\\" not in f]
    files.sort()
    return files


def clean_line(s):
    """移除 NULL 字节和 XML 不兼容的控制字符"""
    return "".join(ch for ch in s if ch == "\t" or ord(ch) >= 32)


def read_all_lines(files):
    all_lines = []
    for f in files:
        rel = os.path.relpath(f, PROJECT_ROOT)
        all_lines.append(f"// ===== {rel} =====")
        try:
            with open(f, "r", encoding="utf-8", errors="replace") as fh:
                for line in fh:
                    all_lines.append(clean_line(line.rstrip("\n")))
        except Exception:
            all_lines.append("// [read error]")
        all_lines.append("")
    # 应用 dream 命名空间替换
    for i, line in enumerate(all_lines):
        for old, new in NAME_REPLACEMENTS:
            line = line.replace(old, new)
        all_lines[i] = line
    return all_lines


def set_cell_font(run, name="Consolas", size=7.5):
    run.font.name = name
    run.font.size = Pt(size)
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), name)


def build_doc(lines, output_path):
    doc = Document()

    # 页面设置 A4
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # 页眉
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hp.add_run(f"{SOFTWARE_NAME}  {VERSION}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 页脚页码
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    fldChar1 = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
    run._element.append(fldChar1)
    run2 = fp.add_run()
    instrText = run2._element.makeelement(qn("w:instrText"), {})
    instrText.text = "PAGE"
    run2._element.append(instrText)
    run3 = fp.add_run()
    fldChar2 = run3._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
    run3._element.append(fldChar2)

    total = len(lines)
    end_front = min(PAGES * LINES_PER_PAGE, total)
    start_back = max(0, total - PAGES * LINES_PER_PAGE)

    def dump_segment(seg_lines, start_global):
        for i, line in enumerate(seg_lines):
            if i > 0 and i % LINES_PER_PAGE == 0:
                doc.add_page_break()
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = Pt(11)
            global_no = start_global + i + 1
            r1 = p.add_run(f"{global_no:>5}  ")
            set_cell_font(r1)
            r1.font.color.rgb = RGBColor(120, 120, 120)
            text = line if len(line) <= 180 else line[:177] + "..."
            r2 = p.add_run(text)
            set_cell_font(r2)

    # 前段
    dump_segment(lines[0:end_front], 0)
    # 分页符分隔前后段
    doc.add_page_break()
    # 后段
    dump_segment(lines[start_back:total], start_back)

    doc.save(output_path)
    print(f"Generated: {output_path}")
    print(f"Total source lines: {total}")
    print(f"Front: lines 1-{end_front} ({end_front} lines, ~{PAGES} pages)")
    print(f"Back:  lines {start_back + 1}-{total} ({total - start_back} lines, ~{PAGES} pages)")


if __name__ == "__main__":
    files = collect_source_files(PROJECT_ROOT)
    print(f"Found {len(files)} source files")
    lines = read_all_lines(files)
    print(f"Total lines (with separators): {len(lines)}")
    build_doc(lines, OUTPUT)
