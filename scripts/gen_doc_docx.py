# -*- coding: utf-8 -*-
"""
生成软件著作权文档鉴别材料 Word 文档
软件设计说明书
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SOFTWARE_NAME = 'One Work 智能协作平台软件'
VERSION = 'V2.1'
OUTPUT = r'D:\aionui-m0\1oneUI\One_Work_软件设计说明书.docx'


def set_font(run, name='微软雅黑', size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    sizes = {1: 16, 2: 14, 3: 12, 4: 11}
    run = p.add_run(text)
    set_font(run, size=sizes.get(level, 11), bold=True)
    return p


def add_para(doc, text, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p


def add_screenshot(doc, filename, caption=None):
    """插入产品截图，居中显示，可选加图注"""
    img_path = os.path.join(r'D:\aionui-m0\1oneUI\resources', filename)
    if not os.path.exists(img_path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(img_path, width=Inches(5.8))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(8)
        crun = cp.add_run(caption)
        set_font(crun, size=9, color=RGBColor(100, 100, 100))


def build_doc():
    doc = Document()

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 页眉
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run(SOFTWARE_NAME + '  ' + VERSION + '  软件设计说明书')
    set_font(run, size=9, color=RGBColor(100, 100, 100))

    # 页脚页码
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    fldChar1 = run._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run._element.append(fldChar1)
    run2 = fp.add_run()
    instrText = run2._element.makeelement(qn('w:instrText'), {})
    instrText.text = 'PAGE'
    run2._element.append(instrText)
    run3 = fp.add_run()
    fldChar2 = run3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run3._element.append(fldChar2)

    # 封面
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(SOFTWARE_NAME)
    set_font(run, size=22, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('（' + VERSION + '）')
    set_font(run, size=16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('软件设计说明书')
    set_font(run, size=20, bold=True)
    for _ in range(8):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('个人独立开发')
    set_font(run, size=14)
    doc.add_page_break()

    # 目录
    add_heading(doc, '目  录', 1)
    toc = [
        '1  引言',
        '  1.1  编写目的',
        '  1.2  项目背景',
        '  1.3  术语定义',
        '  1.4  参考资料',
        '2  总体设计',
        '  2.1  系统架构',
        '  2.2  技术选型',
        '  2.3  运行环境',
        '  2.4  进程模型与通信机制',
        '  2.5  产品核心优势',
        '3  功能模块设计',
        '  3.0  软件主要功能概述',
        '  3.1  会话与对话模块',
        '  3.2  Agent 管理与桥接模块',
        '  3.3  团队协作模块',
        '  3.4  模型管理模块',
        '  3.5  AI 多媒体生成模块',
        '  3.6  定时任务模块',
        '  3.7  MCP 工具模块',
        '  3.8  助手与专家模块',
        '  3.9  记忆管理模块',
        '  3.10  WebUI 远程访问模块',
        '  3.11  IM 渠道接入模块',
        '  3.12  企业管理模块',
        '  3.13  知识库模块',
        '4  接口设计',
        '  4.1  进程间通信',
        '  4.2  HTTP API 接口',
        '  4.3  WebSocket 接口',
        '5  数据设计',
        '  5.1  数据库设计',
        '  5.2  本地文件存储',
        '  5.3  配置数据',
        '6  系统安全设计',
        '  6.1  数据安全',
        '  6.2  访问控制',
        '  6.3  企业级安全',
        '7  部署与运行',
        '  7.1  构建与打包',
        '  7.2  安装与启动',
        '  7.3  升级机制',
        '附录 A  版本历史',
        '附录 B  开发工具与依赖',
        '附录 C  目录结构',
    ]
    for item in toc:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        set_font(run, size=10.5)
    doc.add_page_break()

    # 1 引言
    add_heading(doc, '1  引言', 1)
    add_heading(doc, '1.1  编写目的', 2)
    add_para(doc, '本文档为 One Work 智能协作平台软件（以下简称本软件或 One Work）的软件设计说明书，旨在详细描述本软件的系统架构、功能模块、接口设计、数据结构及安全机制，为软件开发、测试、维护及后续迭代提供技术依据。本文档的预期读者包括软件开发人员、测试人员、运维人员及项目管理人员。')

    add_heading(doc, '1.2  项目背景', 2)
    add_para(doc, '随着人工智能技术的快速发展，AI Agent 在软件开发、企业办公等领域的应用日益广泛。然而，现有 AI 编程工具普遍存在模型厂商锁定、多 Agent 协同能力弱、团队协作支持不足、企业级安全治理缺失等问题。One Work 旨在构建一套开放的 AI Agent 协作平台，支持多模型自由桥接、多 Agent 编队协作、自动化任务执行以及企业级安全治理，帮助个人开发者和企业团队提升研发与办公效率。')
    add_para(doc, '本软件采用前后端分离架构，前端基于 Electron + React 实现跨平台桌面客户端与 WebUI，后端基于 Rust 实现高性能本地服务，支持 Windows、macOS、Linux 三大操作系统。')

    add_heading(doc, '1.3  术语定义', 2)
    terms = [
        ('Agent', '人工智能智能体，能够自主感知环境、做出决策并执行动作的软件实体。'),
        ('CLI', 'Command Line Interface，命令行界面。'),
        ('MCP', 'Model Context Protocol，模型上下文协议，用于连接 AI 模型与外部工具和数据源的开放协议。'),
        ('IPC', 'Inter-Process Communication，进程间通信。'),
        ('RAG', 'Retrieval-Augmented Generation，检索增强生成。'),
        ('SSO', 'Single Sign-On，单点登录。'),
        ('RBAC', 'Role-Based Access Control，基于角色的访问控制。'),
        ('WebUI', '基于浏览器的 Web 用户界面。'),
        ('Team Mode', '团队模式，支持多 Agent 编队协作的运行模式。'),
        ('dreamcore', '本软件后端 Rust 服务的进程名称。'),
    ]
    for term, desc in terms:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.5
        r1 = p.add_run(term + '：')
        set_font(r1, size=10.5, bold=True)
        r2 = p.add_run(desc)
        set_font(r2, size=10.5)

    add_heading(doc, '1.4  参考资料', 2)
    for ref in [
        '《One Work 产品需求文档》',
        '《One Work 用户手册》',
        'Electron 官方文档',
        'React 官方文档',
        'Model Context Protocol 规范',
        'Agent Client Protocol 规范',
    ]:
        add_bullet(doc, ref)
    doc.add_page_break()

    # 2 总体设计
    add_heading(doc, '2  总体设计', 1)
    add_heading(doc, '2.1  系统架构', 2)
    add_para(doc, '本软件采用前后端分离的多层架构，整体分为客户端层、网关层、服务层和数据层四个层次。')
    add_para(doc, '客户端层包括 Electron 桌面客户端和 WebUI 浏览器客户端。Electron 桌面客户端基于 Electron 39 + React 19 + TypeScript 构建，提供完整的桌面应用体验；WebUI 通过独立的 Node 静态服务器托管前端构建产物，支持浏览器和移动设备远程访问。')
    add_para(doc, '网关层由 @dream/web-host 包实现，负责托管 WebUI 静态资源，并将 /api/*、/ws、/login、/logout 等请求反向代理到后端 dreamcore 服务。')
    add_para(doc, '服务层由 dreamcore（Rust 编写的本地服务）实现，包含会话管理、Agent 运行时、模型桥接、MCP 工具、定时任务、企业管理、知识库等 30 余个领域模块。')
    add_para(doc, '数据层采用 SQLite 作为本地数据库，存储会话、助手、消息、企业配置等数据；知识库使用 SQLite FTS5 全文检索与向量检索的混合检索方案。')

    add_heading(doc, '2.2  技术选型', 2)
    add_para(doc, '本软件的主要技术选型如下：')
    for item in [
        '客户端 UI：Electron 39 + React 19 + TypeScript 5.8 + Arco Design',
        '构建工具：Vite 6 + electron-vite 5 + electron-builder',
        '样式方案：UnoCSS + CSS Modules',
        '状态管理：React Hooks + SWR',
        '本地后端：Rust（dreamcore 二进制，30+ 领域 crate）',
        'WebUI 网关：Node.js 静态服务器 + 反向代理',
        '数据库：SQLite（better-sqlite3），FTS5 全文检索',
        '通信协议：HTTP + WebSocket + IPC（Electron 主进程/渲染进程）',
        'AI 协议：Agent Client Protocol (ACP)、Model Context Protocol (MCP)',
        '国际化：i18next，支持 13 种语言',
        '包管理：Bun + npm workspaces',
        '代码质量：Oxlint + Oxfmt + Vitest 4 + Playwright',
    ]:
        add_bullet(doc, item)

    add_heading(doc, '2.3  运行环境', 2)
    add_para(doc, '本软件支持以下运行环境：')
    add_bullet(doc, '操作系统：Windows 10/11（x64/ARM64）、macOS 12 及以上（Intel/Apple Silicon）、Linux（Ubuntu 20.04 及以上，deb 包）')
    add_bullet(doc, '硬件要求：x86_64 或 ARM64 处理器，4GB 及以上内存，200MB 以上可用磁盘空间')
    add_bullet(doc, 'WebUI 浏览器：Chrome、Edge、Safari、Firefox 等现代浏览器')
    add_bullet(doc, 'Node.js 版本：开发环境需 Node.js 22 及以上（运行时由 Electron 内置，用户无需单独安装）')

    add_heading(doc, '2.4  进程模型与通信机制', 2)
    add_para(doc, '本软件采用多进程架构，严格区分主进程（Main Process）与渲染进程（Renderer Process），两者的 API 不可混用。')
    add_para(doc, '主进程运行于 packages/desktop/src/process/ 目录，负责窗口管理、后端进程拉起、系统托盘、自动更新、原生 API 调用等，不可使用 DOM API。')
    add_para(doc, '渲染进程运行于 packages/desktop/src/renderer/ 目录，负责用户界面渲染与交互，不可使用 Node.js API。')
    add_para(doc, '主进程与渲染进程之间通过预加载脚本（preload，packages/desktop/src/preload/）暴露的 IPC 桥接进行通信，确保渲染进程无法直接访问 Node.js 原生能力，符合 Electron 安全最佳实践。')
    add_para(doc, '桌面客户端通过 HTTP + WebSocket 直连 dreamcore 后端服务；WebUI 通过 web-host 网关反向代理访问同一套后端服务，两条路径最终共享同一套数据库与业务逻辑。')

    add_heading(doc, '2.5  产品核心优势', 2)
    add_para(doc, '本软件与市面上通用 AI 聊天工具相比，具备以下四大核心优势：')
    add_heading(doc, '2.5.1  团队协作', 3)
    add_para(doc, '本软件专为团队协作而设计，采用个人版与企业版同一套 UI 的架构。Team Mode 支持多 Agent 编队协作，Leader 拆解任务、Teammate 并行执行；Issues 看板实现产品需求与团队任务的双视图管理，支持 AI 一键拆单；共享会话、组织级 Skills 与 MCP 统一下发，让 PM 提需求、Agent 执行、研发审查的闭环在一个产品内完成。')
    add_heading(doc, '2.5.2  高效交付', 3)
    add_para(doc, '多 Agent 并行执行、Cron 定时任务 7×24 小时无人值守、一句话生成游戏/PPT/文档等成品，配合 WebUI 远程访问与飞书/钉钉/微信/Telegram 等 IM 渠道触达，用户无需频繁切换窗口、等待排期，将时间留给真正重要的决策。内置 22 个官方助手与 252 位行业专家，覆盖办公、研发、金融、法务、设计等领域，开箱即用。')
    add_heading(doc, '2.5.3  简单好用', 3)
    add_para(doc, '内置 Agent 引擎零配置开箱，粘贴模型 API Key 即可开聊；自动检测本机已安装的 38 款主流 CLI Agent 并统一界面管理；20+ 专业助手像同事一样布置任务，非开发者也能轻松使用。支持 13 种语言切换与浅色/深色主题，桌面端、WebUI、手机端随处访问。')
    add_heading(doc, '2.5.4  隐私与信息安全', 3)
    add_para(doc, '对话记录、配置信息与加密后的模型 API Key 全部存储在用户本机磁盘，不经第三方服务器中转；模型请求直连用户选定的供应商，桥接功能为纯本地转发。支持私有化与内网部署，企业版提供 SSO 单点登录、租户隔离、授权许可、模型成本管控与操作/Agent 双维度审计日志，数据不出边界。本软件采用 Apache-2.0 开源协议，免费使用。')
    doc.add_page_break()

    # 3 功能模块
    add_heading(doc, '3  功能模块设计', 1)

    add_heading(doc, '3.0  软件主要功能概述', 2)
    add_para(doc, '本软件是一套 AI Agent 协作平台，核心功能涵盖以下十一个方面：')
    add_para(doc, '一、会话与对话：支持多轮上下文对话，消息类型覆盖文本、图片、文件、代码块与差异对比，Agent 可读写工作区文件并执行多步任务，操作前通过权限对话框请求用户批准，提供 YOLO 全自动、Full-Auto 一键放行与逐次批准三种权限模式，内置 PDF、Word、Excel、PPT、代码、Markdown、图片、HTML 等十余种格式实时预览，无需切换外部软件。')
    add_para(doc, '二、Agent 管理与模型桥接：自动检测本机已安装的 38 款主流 CLI Agent 并统一界面管理，自带内置 Agent 引擎无需额外安装；支持 Codex CLI 与 Claude Code 官方客户端一键桥接至自定义模型服务商，在本机启动兼容协议端点进行纯本地转发，请求全程不出本机，桥接状态在会话顶部实时标识并锁定。')
    add_para(doc, '三、团队协作：Team Mode 支持多 Agent 编队协作，指定 Leader 负责拆解宏观目标并按依赖关系分配子任务，Teammate 并行执行共享工作区，任务看板实时展示进度，成果汇总后支持导出 PDF、HTML、Word 等多种格式。')
    add_para(doc, '四、模型管理：支持 OpenAI、Anthropic、Google、MiniMax、Kimi、DeepSeek 等主流服务商及兼容 OpenAI 协议的自建网关，支持自定义 API 端点、多 API Key 自动轮询与会话内多模型切换。')
    add_para(doc, '五、AI 多媒体生成：集成文生图、图生图、文生视频、图生视频四种生成能力，生成结果标注本次实际花费并支持预估值，视频类任务走异步队列可随时取消以停止计费，应用重启后已提交任务继续回收结果，生成物以卡片形式管理并支持复制、下载与引用回对话。')
    add_para(doc, '六、定时任务：基于 Cron 表达式实现 7×24 小时无人值守自动化，支持新建会话与持续会话两种执行模式，内置队列保护避免任务重叠。')
    add_para(doc, '七、MCP 工具集成：兼容标准 Model Context Protocol，内置浏览器自动化、PDF 导出、金融行情数据等默认连接，支持添加任意自定义 MCP Server，连接测试时逐工具校验 schema 合法性并隔离异常工具。')
    add_para(doc, '八、助手与专家市场：内置 22 个官方助手与 252 位行业专家，覆盖办公文档、研发、金融投研、法务、设计、游戏开发等领域，全部编译期离线分发，断网环境亦可使用，支持一键复制改造或从零自建助手，并支持自定义技能包导入。')
    add_para(doc, '九、记忆管理：采用自动记忆文件、全局配置、项目配置三层记忆架构，与 Claude Code 仓库根目录自动对齐，跨会话沉淀项目上下文。')
    add_para(doc, '十、远程访问与 IM 渠道：WebUI 支持浏览器及移动设备远程访问，具备密码登录、CSRF 防护与速率限制；已接入 Telegram、飞书、钉钉、微信等即时通讯渠道，配对后可直接在聊天窗口指挥 Agent。')
    add_para(doc, '十一、企业级治理：提供个人、项目组、企业三层组织模型，支持 SSO 单点登录、离线授权许可激活、模型成本上限管控、操作与 Agent 双维度审计日志，以及全离线 BM25 加向量混合检索的 RAG 知识库。')
    add_screenshot(doc, '首页.png', '图 3-1  One Work 主界面')
    doc.add_page_break()

    add_heading(doc, '3.1  会话与对话模块', 2)
    add_para(doc, '会话与对话模块是本软件的核心交互模块，提供用户与 AI Agent 之间的多轮对话能力。主要功能包括：')
    add_bullet(doc, '多会话管理：支持创建、切换、删除、重命名会话，所有历史会话集中在会话中心管理')
    add_bullet(doc, '多轮对话：支持上下文关联的多轮对话，消息类型包括文本、图片、文件、代码块、Diff 等')
    add_bullet(doc, '工作区文件操作：Agent 可读写工作区文件、执行多步任务，操作前通过权限对话框请求用户批准')
    add_bullet(doc, '权限控制：支持 YOLO（全自动）、Full-Auto（一键放行）、逐次批准三种权限模式')
    add_bullet(doc, '实时预览：支持 PDF、Word、Excel、PPT、代码、Markdown、图片、HTML、Diff 等 10+ 格式在应用内直接预览')
    add_bullet(doc, '项目面板：三栏布局，右侧实时展示当前项目文件结构与变更，支持搜索定位文件')
    add_para(doc, '该模块前端位于 renderer/pages/ 和 renderer/components/，后端会话管理由 dreamcore 的会话领域 crate 实现，数据持久化到 SQLite。')
    add_screenshot(doc, '工作空间.png', '图 3-2  三栏布局工作区')
    add_screenshot(doc, '历史会话.png', '图 3-3  历史会话中心')

    add_heading(doc, '3.2  Agent 管理与桥接模块', 2)
    add_para(doc, 'Agent 管理与桥接模块负责本机 CLI Agent 的自动检测、统一管理，以及官方 CLI 的模型桥接。主要功能包括：')
    add_bullet(doc, '自动检测：内置支持 38 款主流 AI 编程 Agent（Claude Code、Codex CLI、Cursor、Gemini、Qwen Code 等），已安装的自动识别为可用状态')
    add_bullet(doc, '统一界面：所有 Agent 在统一界面中管理，支持并行多会话')
    add_bullet(doc, '内置引擎：自带 dream CLI Agent 引擎（底层 dreamrs），无需单独安装 CLI，粘贴模型 API Key 即可开聊')
    add_bullet(doc, 'Codex 桥接：在本机启动兼容 OpenAI 协议的本地端点，将 Codex CLI 的请求转发到用户自定义的任意模型服务商（MiniMax、Kimi、DeepSeek、Gemini 等）')
    add_bullet(doc, 'Claude 桥接：同理接管 Claude Code 官方客户端的模型来源，支持一键切换')
    add_bullet(doc, '桥接状态标识：会话顶部模型徽标标注已桥接并锁定，避免实时切模型绕过桥接')
    add_para(doc, '桥接功能为纯本地转发，不经过第三方中转服务器，请求全程不出本机，符合隐私与信息安全优先的设计原则。')
    add_screenshot(doc, 'AGENT助手.png', '图 3-4  Agent 管理界面')
    add_screenshot(doc, 'CODEX和Claude一键桥接.png', '图 3-5  Codex / Claude 官方 CLI 一键桥接')

    add_heading(doc, '3.3  团队协作模块', 2)
    add_para(doc, '团队协作模块支持多 Agent 编队协作，实现 Leader 拆任务、Teammate 并行执行的协同工作模式。主要功能包括：')
    add_bullet(doc, '团队创建：从助手库勾选成员、指定一名 Leader 即可组建团队')
    add_bullet(doc, '任务拆解：Leader Agent 负责将宏观目标拆解为子任务，按依赖关系分配给 Teammate')
    add_bullet(doc, '并行执行：多个 Teammate Agent 并行执行各自任务，共享工作区')
    add_bullet(doc, '任务看板：共享任务看板，实时展示各任务状态与进度')
    add_bullet(doc, '成果汇总：Leader 汇总各 Teammate 的产出，生成最终交付物（支持导出 PDF/HTML/Word 等格式）')
    add_bullet(doc, '动态组队：针对复杂目标可现场组建团队，灵活配置成员与角色')
    add_para(doc, '该模块基于 Agent Client Protocol（ACP）实现多 Agent 间的通信与协调，后端由 dreamcore 的 team 领域 crate 提供支持。')
    add_screenshot(doc, '团队多任务角色.png', '图 3-6  创建团队：选择成员与 Leader')
    add_screenshot(doc, '团队AGENT测试.png', '图 3-7  Team Mode 多 Agent 协作过程')

    add_heading(doc, '3.4  模型管理模块', 2)
    add_para(doc, '模型管理模块负责 AI 模型服务商与模型的配置、管理与调度。主要功能包括：')
    add_bullet(doc, '多平台支持：支持 OpenAI、Anthropic、Google、MiniMax、月之暗面（Kimi）、DeepSeek、阿里云、腾讯云、火山引擎等主流模型服务商，以及兼容 OpenAI 协议的自定义网关')
    add_bullet(doc, '自定义 API 地址：支持填写自定义 API 端点，适配自建模型网关')
    add_bullet(doc, '多 API Key 轮询：同一模型平台可配置多个 API Key，自动轮询使用，提升稳定性')
    add_bullet(doc, '多模型选择：一次勾选多个模型，会话中可随时切换')
    add_bullet(doc, '模型类型标注：支持将模型标注为文本模型、图片模型或视频模型，自动出现在对应生成模式的可选项中')
    add_bullet(doc, '加密存储：模型 API Key 加密后存储在本地，密钥字段永不回显')
    add_screenshot(doc, '模型添加.png', '图 3-8  模型平台配置')

    add_heading(doc, '3.5  AI 多媒体生成模块', 2)
    add_para(doc, 'AI 多媒体生成模块集成文生图、图生图、文生视频、图生视频四种生成能力。主要功能包括：')
    add_bullet(doc, '文生图：输入提示词生成图片，支持一次生成最多 9 张互相独立的图片')
    add_bullet(doc, '图生图：以参考图为基础生成新图片，结果卡片同时显示参考图与产出图')
    add_bullet(doc, '文生视频：输入提示词生成视频，完成后内联播放')
    add_bullet(doc, '图生视频：以图片为首帧生成动态视频')
    add_bullet(doc, '花费透明：生成结果卡片标注本次实际花费，发送前给出预估费用，读不到真实价格时明示费用未知')
    add_bullet(doc, '异步任务与取消：视频生成走异步任务，支持随时取消（取消即停止计费），应用重启后已提交任务继续回收结果')
    add_bullet(doc, '结果管理：生成结果以卡片形式展示，支持复制、下载、引用回输入框继续追问、重新生成')
    add_bullet(doc, '专家调用：图片/视频生成作为内置 MCP 工具，助手和行业专家可自主调用')
    add_para(doc, '生成物存储在用户选择的工作目录中，与会话其他产物放在一起。')
    add_screenshot(doc, '文生图.png', '图 3-9  文生图结果卡片')
    add_screenshot(doc, '文生视频.png', '图 3-10  文生视频内联播放')

    add_heading(doc, '3.6  定时任务模块', 2)
    add_para(doc, '定时任务模块基于 Cron 表达式实现 7×24 小时自动化任务执行。主要功能包括：')
    add_bullet(doc, 'Cron 调度：支持标准 Cron 表达式配置执行频率')
    add_bullet(doc, '执行助手绑定：每个定时任务可绑定任意官方助手或行业专家作为执行者')
    add_bullet(doc, '两种会话模式：新建会话模式（每次干净上下文，适合日报巡检）与持续会话模式（沿用上下文，适合长期跟进）')
    add_bullet(doc, '队列保护：可开启队列保护，避免上一次任务还在运行就被新触发打断')
    add_bullet(doc, '任务管理：支持查看任务执行历史、手动触发、启用/禁用、删除')
    add_para(doc, '该模块基于 croner 库实现调度逻辑，任务执行通过 Agent 运行时完成。')
    add_screenshot(doc, '定时任务.png', '图 3-11  定时任务配置')

    add_heading(doc, '3.7  MCP 工具模块', 2)
    add_para(doc, 'MCP 工具模块实现标准 Model Context Protocol，用于连接外部工具与数据源。主要功能包括：')
    add_bullet(doc, '标准 MCP 协议：完全兼容 MCP 规范，可接入任意标准 MCP Server')
    add_bullet(doc, '内置默认连接：随包内置浏览器自动化（chrome-devtools）、团队检索（team-search）、PDF 导出（one-export-pdf）、金融行情数据（ftshare，覆盖 A 股/港股/美股 170+ 工具）等')
    add_bullet(doc, '自定义 MCP Server：支持自行添加任意 MCP Server，配置传输方式（stdio/SSE/HTTP）')
    add_bullet(doc, '逐工具 schema 校验：连接测试时逐个工具校验 schema 合法性，只隔离有问题的工具而不牵连整个 Server')
    add_bullet(doc, '配置隔离：MCP 配置读写与用户系统中 Claude Code/Codex 的配置完全隔离，互不干扰')
    add_screenshot(doc, 'MCP工具.png', '图 3-12  MCP 工具配置')

    add_heading(doc, '3.8  助手与专家模块', 2)
    add_para(doc, '助手与专家模块提供开箱即用的角色库，分为我的助手、官方助手、专家市场三个层级。主要功能包括：')
    add_bullet(doc, '官方助手（22 个）：覆盖 Word、PPT、Excel、可填表单、数据仪表盘、财务建模、学术论文、3D 游戏生成、故事角色扮演等场景，全部配好技能与工具')
    add_bullet(doc, '专家市场（252 位）：覆盖云运维、金融投研、法务、销售、设计、游戏开发、公益、医疗合规等各行业岗位，带真实中文名与头像，一键添加即用')
    add_bullet(doc, '离线分发：252 位专家人设在编译期打进后端二进制，随安装包分发，不联网、不额外下载')
    add_bullet(doc, '自建助手：支持一键复制官方助手改造，或从零组合 Agent + 技能 + 规则创建自定义助手')
    add_bullet(doc, '技能库：20+ 内置技能覆盖社交招聘发帖、故事角色扮演、微信文件回传、OpenClaw 部署、PDF 处理、Office 文档等场景；支持导入自定义技能（目录或 zip，单文件 50MB、单技能 200MB）')
    add_bullet(doc, '搜索合并：搜索时自动合并已安装与市场目录，点未安装条目自动安装并选中')
    add_screenshot(doc, '专家市场.png', '图 3-13  助手与专家市场')
    add_screenshot(doc, '内置技能.png', '图 3-14  内置技能库')

    add_heading(doc, '3.9  记忆管理模块', 2)
    add_para(doc, '记忆管理模块实现三层记忆架构，让 AI 越用越懂用户的项目。三层记忆架构包括：')
    add_bullet(doc, '自动记忆（MEMORY.md）：Agent 在对话过程中自动沉淀的关键信息，跨会话持久化')
    add_bullet(doc, '全局 CLAUDE.md：全局级别的项目背景与角色定位，适用于所有项目')
    add_bullet(doc, '项目 CLAUDE.md：项目级别的上下文信息，与 Claude Code 仓库根目录自动对齐')
    add_para(doc, '记忆文件存储在工作区中，Agent 可自动读取和更新，用户也可手动编辑。')
    add_screenshot(doc, '记忆.png', '图 3-15  三层记忆管理')

    add_heading(doc, '3.10  WebUI 远程访问模块', 2)
    add_para(doc, 'WebUI 远程访问模块允许用户通过浏览器或移动设备远程访问本软件。主要功能包括：')
    add_bullet(doc, 'HTTP 服务：基于 Express 5 启动本地 HTTP 服务器，托管 WebUI 静态资源')
    add_bullet(doc, '反向代理：将 /api/*、/ws、/login、/logout 等请求代理到 dreamcore 后端')
    add_bullet(doc, '身份认证：支持密码登录和扫码登录，保障远程访问安全')
    add_bullet(doc, '权限控制：可配置 WebUI 账号权限，限制远程用户的操作范围')
    add_bullet(doc, '局域网访问：支持局域网内其他设备访问，手机扫码即可连接')
    add_bullet(doc, 'CSRF 防护：使用 tiny-csrf 中间件防护跨站请求伪造')
    add_bullet(doc, '速率限制：使用 express-rate-limit 限制请求频率，防止暴力攻击')
    add_screenshot(doc, '远程访问.png', '图 3-16  WebUI 远程访问')

    add_heading(doc, '3.11  IM 渠道接入模块', 2)
    add_para(doc, 'IM 渠道接入模块将本软件能力延伸到即时通讯软件，用户可直接在聊天窗口中指挥 Agent。已支持渠道包括：')
    add_bullet(doc, 'Telegram：基于 grammy 框架实现')
    add_bullet(doc, '飞书（Lark）：基于 @larksuiteoapi/node-sdk 实现')
    add_bullet(doc, '钉钉：基于 dingtalk-stream 实现')
    add_bullet(doc, '微信：基于 @wecom/aibot-node-sdk 实现')
    add_para(doc, '规划中渠道包括企业微信、Slack、Discord。每个渠道独立配置，配对成功后即可在对应 IM 中发起对话、执行任务、接收结果。')
    add_screenshot(doc, '渠道配置.png', '图 3-17  IM 渠道配置')

    add_heading(doc, '3.12  企业管理模块', 2)
    add_para(doc, '企业管理模块在个人版功能基础上叠加治理层，采用个人、项目组、企业三层模型。')
    add_para(doc, '项目组功能包括：')
    add_bullet(doc, '局域网创建：无需实名、无需公网，局域网内创建项目组，邀请码拉人')
    add_bullet(doc, '服务器/客户端模式：同一时间一台机器作为服务器托管数据，其余自动连接，可随时切换角色')
    add_bullet(doc, '团队管理：成员与角色管理、共享技能/MCP/知识库/流水线、组织架构')
    add_bullet(doc, '研发智能工作台：协作看板、数字员工、流水线编排、团队知识库、版本规划、测试计划、运行时节点、效能洞察、制品仓库')
    add_bullet(doc, '双维度审计：操作审计与 Agent 执行审计')
    add_para(doc, '企业功能包括：')
    add_bullet(doc, '企业身份设立：管理员显式设立企业身份')
    add_bullet(doc, 'SSO 单点登录：支持飞书、钉钉、企业微信、LDAP 及标准 OIDC（适配 Okta、Azure AD、Google Workspace 等）')
    add_bullet(doc, '租户隔离：企业间数据完全隔离')
    add_bullet(doc, '授权许可：离线激活，粘贴厂商签发的授权码即可升档，签名校验为唯一升档路径')
    add_bullet(doc, '订阅与席位：free/team/enterprise 三档，功能按档位门控，席位用量实时可见')
    add_bullet(doc, '模型成本管控：设置近 30 天成本上限与模型 allowlist，从源头控制开销')
    add_bullet(doc, '备份与恢复：企业数据整库导出/导入，导出时自动脱敏密钥字段')
    add_bullet(doc, '需求协作看板（Issues）：产品需求与团队任务双视图，看板拖拽流转，支持 AI Agent 需求一键拆单')
    add_bullet(doc, '超级助手：Issues 与 Agent 执行打通，从看板选中任务一键发起 Agent 处理，管理多智能体并行执行')
    add_bullet(doc, '数字员工：可绑定任意专家人设 + 运行后端 + 具体模型，按计划自动执行并留下运行历史')
    add_screenshot(doc, '项目组总后台.png', '图 3-18  项目组管理后台')
    add_screenshot(doc, '企业版后台.png', '图 3-19  企业版管理后台')

    add_heading(doc, '3.13  知识库模块', 2)
    add_para(doc, '知识库模块实现全离线的语义 + 词法混合检索，为 Agent 提供团队知识检索能力。主要功能包括：')
    add_bullet(doc, '混合检索：BM25 词法检索与向量语义检索的 RRF（Reciprocal Rank Fusion）融合')
    add_bullet(doc, '中文分词：内置中文分词支持，适配中文文档检索')
    add_bullet(doc, '全离线运行：不依赖外部向量数据库或云服务，数据全部本地存储')
    add_bullet(doc, 'MCP 工具接入：知识库以 MCP 工具形式接入 Agent，Agent 可自主检索相关文档')
    add_bullet(doc, '权限复用：知识库权限自动复用团队 ACL，确保数据安全')
    add_bullet(doc, '文档管理：支持上传、分类、检索团队文档')
    doc.add_page_break()

    # 4 接口设计
    add_heading(doc, '4  接口设计', 1)
    add_heading(doc, '4.1  进程间通信', 2)
    add_para(doc, 'Electron 主进程与渲染进程之间通过预加载脚本暴露的 IPC 桥接进行通信。预加载脚本位于 packages/desktop/src/preload/，使用 contextBridge.exposeInMainWorld 将安全的 API 暴露给渲染进程。')
    add_para(doc, 'IPC 接口按功能域划分，主要包括：')
    add_bullet(doc, '窗口管理：窗口创建、最小化、最大化、关闭、置顶、全屏等')
    add_bullet(doc, '文件操作：文件选择对话框、文件读写、目录选择、拖放文件处理')
    add_bullet(doc, '系统托盘：托盘图标、菜单、通知、气泡提示')
    add_bullet(doc, '自动更新：检查更新、下载更新、安装更新、更新进度回调')
    add_bullet(doc, '后端进程管理：dreamcore 进程启动、停止、重启、健康检查、日志收集')
    add_bullet(doc, '原生能力：剪贴板操作、全局快捷键、系统信息获取、电源管理')
    add_para(doc, '所有 IPC 通道在主进程端注册监听器，渲染进程通过 window.api 调用，严格遵循请求-响应模式，支持异步回调。')

    add_heading(doc, '4.2  HTTP API 接口', 2)
    add_para(doc, 'dreamcore 后端提供 RESTful HTTP API，供桌面客户端和 WebUI 调用。API 统一前缀为 /api/v1/，采用 JSON 格式交换数据。主要 API 分组包括：')
    add_bullet(doc, '/api/v1/auth：登录、登出、令牌刷新、会话验证')
    add_bullet(doc, '/api/v1/conversations：会话创建、列表、详情、更新、删除、消息发送')
    add_bullet(doc, '/api/v1/agents：Agent 列表、状态检测、配置管理')
    add_bullet(doc, '/api/v1/models：模型平台配置、模型列表、模型测试')
    add_bullet(doc, '/api/v1/assistants：助手列表、创建、更新、删除、专家市场同步')
    add_bullet(doc, '/api/v1/mcp：MCP 连接管理、工具列表、工具调用测试')
    add_bullet(doc, '/api/v1/cron：定时任务 CRUD、手动触发、执行历史')
    add_bullet(doc, '/api/v1/media：图片/视频生成任务提交、状态查询、取消、结果获取')
    add_bullet(doc, '/api/v1/teams：团队管理、成员管理、任务分配、执行状态')
    add_bullet(doc, '/api/v1/enterprise：企业配置、成员管理、SSO 配置、授权许可、审计日志、备份恢复')
    add_bullet(doc, '/api/v1/knowledge：知识库文档管理、检索接口')
    add_bullet(doc, '/api/v1/channels：IM 渠道配置、配对、状态管理')
    add_para(doc, 'API 认证采用 Bearer Token 机制，登录后获取访问令牌，后续请求在 Authorization 头中携带。')

    add_heading(doc, '4.3  WebSocket 接口', 2)
    add_para(doc, '对于需要实时推送的场景（如 Agent 执行过程中的流式输出、任务进度更新、实时日志），本软件采用 WebSocket 通信。')
    add_para(doc, 'WebSocket 端点为 /ws，连接建立后通过 JSON 消息进行双向通信。主要消息类型包括：')
    add_bullet(doc, 'conversation.message：对话消息的流式增量推送（Token 级实时输出）')
    add_bullet(doc, 'agent.status：Agent 运行状态变更通知')
    add_bullet(doc, 'task.progress：定时任务、生成任务的进度更新')
    add_bullet(doc, 'file.change：工作区文件变更实时通知')
    add_bullet(doc, 'team.event：团队协作事件（任务分配、成员状态、产出更新）')
    add_bullet(doc, 'system.notification：系统级通知推送')
    add_para(doc, 'WebSocket 连接支持心跳保活，断线自动重连，确保实时消息不丢失。')
    doc.add_page_break()

    # 5 数据设计
    add_heading(doc, '5  数据设计', 1)
    add_heading(doc, '5.1  数据库设计', 2)
    add_para(doc, '本软件采用 SQLite 作为本地数据库，数据库文件存储在应用数据目录下（Windows 为 %APPDATA%，macOS 为 ~/Library/Application Support，Linux 为 ~/.config），按环境（开发/生产）隔离。主要数据表包括：')
    for table in [
        'conversations：会话表，存储会话 ID、标题、创建时间、更新时间、绑定助手、模型配置等',
        'messages：消息表，存储消息 ID、会话 ID、角色（user/assistant/tool）、内容、时间戳、元数据',
        'agents：Agent 配置表，存储 Agent 类型、安装路径、版本、启用状态、配置参数',
        'models：模型平台配置表，存储平台名称、API 地址、API Key（加密）、模型列表、轮询配置',
        'assistants：助手表，存储助手 ID、名称、描述、人设、绑定 Agent、技能列表、规则配置',
        'experts：专家市场表，存储专家 ID、名称、行业、头像、人设、技能、安装状态',
        'mcp_servers：MCP 服务器配置表，存储名称、传输方式、命令/URL、环境变量、连接状态',
        'mcp_tools：MCP 工具表，存储所属服务器、工具名称、描述、参数 Schema、校验状态',
        'cron_jobs：定时任务表，存储任务 ID、名称、Cron 表达式、绑定助手、执行指令、会话模式、启用状态、队列保护配置',
        'cron_runs：定时任务执行历史表，存储任务 ID、开始时间、结束时间、状态、输出摘要',
        'media_tasks：多媒体生成任务表，存储任务 ID、类型（图/视频）、提示词、参考图、模型、状态、花费、结果路径',
        'teams：团队表，存储团队 ID、名称、Leader、成员列表、创建时间',
        'team_tasks：团队任务表，存储任务 ID、团队 ID、分配者、执行者、状态、依赖关系、产出',
        'organizations：企业/项目组表，存储组织 ID、名称、类型（项目组/企业）、部署模式、创建者',
        'members：组织成员表，存储组织 ID、用户 ID、角色、加入时间、状态',
        'audit_logs：审计日志表，存储操作人、操作类型、目标对象、时间、IP、详情',
        'knowledge_docs：知识库文档表，存储文档 ID、组织 ID、标题、内容、向量、标签、上传时间',
        'channels：IM 渠道配置表，存储渠道类型、配置参数、配对状态、绑定助手',
    ]:
        add_bullet(doc, table)
    add_para(doc, '数据库使用 better-sqlite3 进行同步访问，配合 WAL 模式提升并发性能，所有写操作使用事务保证原子性。')

    add_heading(doc, '5.2  本地文件存储', 2)
    add_para(doc, '除数据库外，本软件还在本地文件系统中存储以下数据：')
    add_bullet(doc, '会话工作区：用户选择的工作目录，Agent 产生的文件、生成的图片/视频、导出的文档等均存储于此')
    add_bullet(doc, '记忆文件：MEMORY.md、全局 CLAUDE.md、项目 CLAUDE.md，存储在对应工作区根目录')
    add_bullet(doc, '技能文件：内置技能随安装包分发，自定义技能导入后存储在应用数据目录的 skills 子目录')
    add_bullet(doc, '日志文件：应用日志、dreamcore 后端日志，存储在 logs 子目录，按日期轮转')
    add_bullet(doc, '配置文件：应用配置、主题设置、语言设置等，存储在 config 子目录')
    add_bullet(doc, '缓存文件：图片缩略图、文档预览缓存、模型图标等，存储在 cache 子目录')

    add_heading(doc, '5.3  配置数据', 2)
    add_para(doc, '应用配置采用 JSON 格式存储，主要配置项包括：')
    add_bullet(doc, '通用设置：语言、主题、主题色、字体大小、窗口状态')
    add_bullet(doc, '代理设置：HTTP 代理、SOCKS 代理、代理白名单')
    add_bullet(doc, '安全设置：API Key 加密密钥、会话锁定、自动锁定时间')
    add_bullet(doc, 'WebUI 设置：服务端口、启用状态、密码、远程访问权限')
    add_bullet(doc, '桥接设置：Codex 桥接开关与配置、Claude 桥接开关与配置')
    add_bullet(doc, 'Agent 设置：Agent 扫描路径、默认 Agent、权限模式')
    add_bullet(doc, '更新设置：自动检查更新、更新通道、代理设置')
    add_para(doc, '敏感配置（如 API Key、密码）使用 bcryptjs 进行哈希或 AES 加密存储，绝不以明文形式写入配置文件。')
    doc.add_page_break()

    # 6 系统安全
    add_heading(doc, '6  系统安全设计', 1)
    add_heading(doc, '6.1  数据安全', 2)
    add_para(doc, '本软件遵循隐私与信息安全优先的设计原则，所有用户数据默认存储在本地，不上传第三方服务器。')
    add_bullet(doc, '本地存储：会话、配置、加密后的模型 API Key 全部落在本机磁盘，不外传')
    add_bullet(doc, 'API Key 加密：模型 API Key 使用 AES 加密存储，密钥由用户本机生成，永不回显')
    add_bullet(doc, '密码哈希：WebUI 登录密码、企业管理员密码使用 bcryptjs 加盐哈希存储')
    add_bullet(doc, 'JWT 认证：API 访问使用 jsonwebtoken 签发令牌，设置过期时间，支持刷新')
    add_bullet(doc, '桥接本地转发：Codex/Claude 桥接为纯本地转发，不经过第三方中转服务器')
    add_bullet(doc, '备份脱敏：企业数据导出时自动脱敏密钥字段（如飞书 appSecret 被剥离）')

    add_heading(doc, '6.2  访问控制', 2)
    add_bullet(doc, 'Electron 安全：启用 contextIsolation、禁用 nodeIntegration、启用 sandbox，渲染进程无法直接访问 Node.js API')
    add_bullet(doc, 'IPC 白名单：预加载脚本只暴露必要的安全 API，渲染进程无法直接调用任意主进程功能')
    add_bullet(doc, 'WebUI 认证：远程访问需密码或扫码登录，支持配置账号权限')
    add_bullet(doc, 'CSRF 防护：WebUI 使用 tiny-csrf 中间件防护跨站请求伪造')
    add_bullet(doc, '速率限制：使用 express-rate-limit 限制 API 请求频率，防止暴力破解与滥用')
    add_bullet(doc, 'CORS 配置：WebUI 网关配置 cors 中间件，限制跨域访问来源')
    add_bullet(doc, 'Agent 权限：Agent 执行文件操作前需用户批准，支持 YOLO/Full-Auto/逐次批准三种模式')

    add_heading(doc, '6.3  企业级安全', 2)
    add_bullet(doc, '租户隔离：企业间数据完全隔离，每个企业独立数据库作用域')
    add_bullet(doc, 'SSO 单点登录：支持飞书、钉钉、企业微信、LDAP、OIDC 等标准身份提供商')
    add_bullet(doc, 'RBAC 权限：基于角色的细粒度访问控制，功能按档位门控')
    add_bullet(doc, '操作审计：操作与 Agent 执行双维度审计日志，可追溯所有关键操作')
    add_bullet(doc, '离职失效：移除成员后即时失效其会话，并归因到操作者')
    add_bullet(doc, '授权许可：离线激活，签名校验为唯一升档路径，管理员无法自行改档')
    add_bullet(doc, '模型成本管控：成本上限与模型 allowlist，防止失控的模型开销')
    add_bullet(doc, '备份恢复：企业数据整库导出/导入，往返幂等，支持灾难恢复')
    doc.add_page_break()

    # 7 部署与运行
    add_heading(doc, '7  部署与运行', 1)
    add_heading(doc, '7.1  构建与打包', 2)
    add_para(doc, '本软件使用 electron-vite 进行构建，electron-builder 进行打包发行。构建流程如下：')
    add_bullet(doc, '前端构建：Vite 将 React 渲染进程代码打包为静态资源，输出到 out/renderer/')
    add_bullet(doc, '主进程构建：Vite 将主进程 TypeScript 代码打包为 CommonJS，输出到 out/main/')
    add_bullet(doc, '后端内嵌：将编译好的 dreamcore 二进制（Rust）复制到 resources 目录，随安装包分发')
    add_bullet(doc, '打包：electron-builder 根据目标平台生成对应安装包（Windows exe/msi、macOS dmg、Linux deb）')
    add_para(doc, '支持的构建命令包括：')
    add_bullet(doc, 'bun run dist:win：构建 Windows 安装包（x64/ARM64）')
    add_bullet(doc, 'bun run dist:mac：构建 macOS 安装包（Intel/Apple Silicon）')
    add_bullet(doc, 'bun run dist:linux：构建 Linux deb 安装包')
    add_bullet(doc, 'bun run build：构建全平台安装包')

    add_heading(doc, '7.2  安装与启动', 2)
    add_para(doc, '用户安装流程：')
    add_bullet(doc, 'Windows：下载 .exe 安装包，双击运行，按向导完成安装，支持 Squirrel 安装器')
    add_bullet(doc, 'macOS：下载 .dmg 镜像，双击打开，将应用拖入 Applications 文件夹')
    add_bullet(doc, 'Linux：下载 .deb 包，使用 dpkg -i 或软件中心安装')
    add_para(doc, '启动流程：')
    add_bullet(doc, 'Electron 主进程启动，初始化窗口、系统托盘、日志系统')
    add_bullet(doc, '主进程拉起 bundled 的 dreamcore 后端进程，等待后端健康检查通过')
    add_bullet(doc, '加载渲染进程页面，建立 IPC 连接与 HTTP/WebSocket 连接')
    add_bullet(doc, '检查更新（如启用自动更新），加载用户配置与历史会话')
    add_bullet(doc, '显示主窗口，进入应用首页')

    add_heading(doc, '7.3  升级机制', 2)
    add_para(doc, '本软件使用 electron-updater 实现自动更新，支持增量更新与全量更新。')
    add_bullet(doc, '更新检查：启动时自动检查更新，或用户手动触发检查')
    add_bullet(doc, '更新下载：发现新版本后后台下载安装包，显示下载进度')
    add_bullet(doc, '更新安装：下载完成后提示用户重启安装，或在退出时自动安装')
    add_bullet(doc, '更新源：支持 GitHub Releases、自定义更新服务器，可配置代理')
    add_bullet(doc, '版本回退：安装包保留历史版本，用户可手动降级')
    add_bullet(doc, '后端同步更新：dreamcore 后端随前端安装包一起更新，确保版本兼容')

    # 附录
    doc.add_page_break()
    add_heading(doc, '附录 A  版本历史', 1)
    add_para(doc, '本文档对应软件版本 V2.1，主要版本演进如下：', indent=False)
    add_bullet(doc, 'V1.x：初始版本，单仓架构（One Work ClaudeCode），基础 AI 对话与 Claude Code 集成')
    add_bullet(doc, 'V2.0：前后端分离重构（dream + dreamcore），引入 Electron + React + Rust 架构，支持多 Agent、WebUI、IM 渠道')
    add_bullet(doc, 'V2.1：新增企业版（项目组/企业三层模型）、SSO、授权许可、知识库 RAG、需求协作看板、超级助手、AI 多媒体生成等能力')

    add_heading(doc, '附录 B  开发工具与依赖', 1)
    add_para(doc, '本软件主要开发依赖与版本：', indent=False)
    for dep in [
        '运行时：Electron 39.8.10、Node.js 22+、Bun',
        '前端框架：React 19.1.0、React Router 7.8.0、TypeScript 5.8.3',
        'UI 组件库：@arco-design/web-react 2.66.1、@icon-park/react 1.4.2',
        '构建工具：Vite 6.4.1、electron-vite 5.0.0、electron-builder 26.15.2',
        '样式：UnoCSS 66.3.3、PostCSS 8.5.8',
        '代码编辑器：@uiw/react-codemirror 4.25.2、@monaco-editor/react 4.7.0',
        'Markdown：react-markdown 10.1.0、remark-gfm 4.0.1、rehype-katex 7.0.1',
        'AI SDK：@anthropic-ai/sdk 0.71.2、openai 5.12.2、@google/genai 1.16.0、@modelcontextprotocol/sdk 1.20.0',
        '数据库：better-sqlite3 12.4.1',
        'WebUI 服务：express 5.1.0、ws 8.18.3、cors 2.8.5、cookie-parser 1.4.7',
        '安全：bcryptjs 2.4.3、jsonwebtoken 9.0.2、tiny-csrf 1.1.6、express-rate-limit 7.5.1',
        '文档处理：docx 9.5.1、mammoth 1.11.0、xlsx-republish 0.20.3、pptx2json 0.0.10、officeparser 5.2.2',
        '测试：Vitest 4.0.18、Playwright 1.59.1、@testing-library/react 16.3.2',
        '代码质量：Oxlint 1.56.0、Oxfmt 0.41.0、TypeScript 5.8.3',
    ]:
        add_bullet(doc, dep)

    add_heading(doc, '附录 C  目录结构', 1)
    add_para(doc, '项目采用 monorepo 结构，主要目录如下：', indent=False)
    for d in [
        'packages/desktop/：Electron 桌面客户端主包，包含主进程、渲染进程、预加载脚本',
        'packages/desktop/src/process/：主进程代码（窗口管理、后端进程、系统托盘、自动更新）',
        'packages/desktop/src/renderer/：渲染进程代码（React UI、页面、组件、服务、Hooks）',
        'packages/desktop/src/preload/：预加载脚本（IPC 桥接）',
        'packages/desktop/src/common/：公共代码（i18n 配置、常量、类型）',
        'packages/web-host/：WebUI 静态服务器与反向代理',
        'packages/web-cli/：Web CLI 工具',
        'packages/shared-scripts/：共享脚本',
        'resources/：应用资源（图标、截图、bundled 后端）',
        'scripts/：构建、开发、工具脚本',
        'docs/：项目文档',
        'tests/：测试用例（单元测试、集成测试、E2E 测试、性能基准）',
        'locales/：国际化语言包（13 种语言）',
    ]:
        add_bullet(doc, d)

    doc.save(OUTPUT)
    print('Generated:', OUTPUT)


if __name__ == '__main__':
    build_doc()
